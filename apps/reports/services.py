from pathlib import Path

from django.conf import settings
from django.core.exceptions import PermissionDenied, ValidationError
from django.db import transaction
from django.utils import timezone
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from apps.approvals.services import revoke_active_signatures, sign_approval
from apps.audit.services import record_audit, snapshot

from .models import RegulatoryReport

REPORT_STATUS_TRANSITIONS = {
    RegulatoryReport.Status.DRAFT: {RegulatoryReport.Status.REVIEW_PENDING},
    RegulatoryReport.Status.REVIEW_PENDING: {
        RegulatoryReport.Status.APPROVED,
        RegulatoryReport.Status.REJECTED,
    },
    RegulatoryReport.Status.REJECTED: {RegulatoryReport.Status.REVIEW_PENDING},
    RegulatoryReport.Status.APPROVED: {
        RegulatoryReport.Status.GENERATED,
        RegulatoryReport.Status.SUBMITTED,
    },
    RegulatoryReport.Status.GENERATED: {
        RegulatoryReport.Status.GENERATED,
        RegulatoryReport.Status.SUBMITTED,
    },
    RegulatoryReport.Status.SUBMITTED: set(),
}


def visible_reports(user):
    qs = RegulatoryReport.objects.all()
    return qs.filter(adverse_event__reporter=user) if user.role == "STAFF" else qs


def _validate_status_transition(report, new_status):
    if new_status not in RegulatoryReport.Status.values:
        raise ValidationError("알 수 없는 보고서 상태입니다.")
    if new_status not in REPORT_STATUS_TRANSITIONS[report.report_status]:
        raise ValidationError(
            f"보고서 상태를 {report.report_status}에서 {new_status}(으)로 변경할 수 없습니다."
        )


def _allowed(user):
    if user.role not in {"RA_QA","ADMIN"}: raise PermissionDenied("RA·QA 또는 ADMIN만 보고서를 변경할 수 있습니다.")
def _audit(user,action,report,before=None,request=None):
    record_audit(user=user,action=action,target=report,before=before or {},after=snapshot(report),request=request)
def generate_report_number():
    from apps.compliance.services import next_management_number
    return next_management_number("REGULATORY_REPORT")


@transaction.atomic
def reject_report(report,user,password=None,reason="",request=None):
    if user.role!="ADMIN": raise PermissionDenied("ADMIN만 보고서를 반려할 수 있습니다.")
    _validate_status_transition(report,RegulatoryReport.Status.REJECTED)
    sign_approval(target=report,user=user,password=password,meaning="규제보고서 반려",reason=reason,request=request,forbid_self_user_id=report.created_by_id)
    before=snapshot(report); report.report_status=RegulatoryReport.Status.REJECTED; report.approved_by=None; report.save(); record_audit(user=user,action="REPORT_REJECT",target=report,before=before,after=snapshot(report),reason=reason,request=request,require_reason=True); return report


@transaction.atomic
def cancel_report_approval(report,user,password=None,reason="",request=None):
    if user.role!="ADMIN": raise PermissionDenied("ADMIN만 승인을 취소할 수 있습니다.")
    if report.report_status not in {RegulatoryReport.Status.APPROVED,RegulatoryReport.Status.GENERATED}: raise ValidationError("승인 또는 생성 상태의 보고서만 승인 취소할 수 있습니다.")
    if not password or not user.check_password(password): raise ValidationError("승인자 재인증에 실패했습니다.")
    if not reason.strip(): raise ValidationError("승인 취소 사유가 필요합니다.")
    before=snapshot(report); revoke_active_signatures(report,user=user,reason=reason,request=request); report.report_status=RegulatoryReport.Status.DRAFT; report.approved_by=None; report.save(); record_audit(user=user,action="REPORT_APPROVAL_CANCEL",target=report,before=before,after=snapshot(report),reason=reason,request=request,require_reason=True); return report
def populate_report_fields(event):
    patient=getattr(event,"patient_info",None); investigation=getattr(event,"investigation",None)
    return {"title":f"{event.event_number} 규제 보고서","event_summary":f"{event.title}\n{event.description}","device_information":f"제품: {event.medical_device.product_name}\n모델: {event.medical_device.model_name}\n허가번호: {event.medical_device.approval_number}\nLOT: {event.device_lot.lot_number}\n시리얼: {event.device_lot.serial_number or '미입력'}","patient_information":f"익명코드: {patient.anonymous_code}\n연령군: {patient.age_group}\n성별: {patient.gender}" if patient else "미입력","investigation_summary":investigation.investigation_summary if investigation else "","root_cause_summary":investigation.root_cause if investigation else "","capa_summary":"\n".join(f"{c.capa_number}: {c.corrective_action} / {c.preventive_action}" for c in event.capas.all())}
def validate_report_generation(report):
    if not report.adverse_event.medical_device_id: raise ValidationError("제품 정보가 필요합니다.")
    if report.report_status not in {RegulatoryReport.Status.APPROVED,RegulatoryReport.Status.GENERATED}: raise ValidationError("승인된 보고서만 DOCX를 생성할 수 있습니다.")
@transaction.atomic
def create_report_from_event(event,user,**overrides):
    _allowed(user); data=populate_report_fields(event); data.update(overrides); report=RegulatoryReport.objects.create(adverse_event=event,created_by=user,**data); _audit(user,"REPORT_CREATE",report); return report
def request_report_review(report,user):
    _allowed(user); _validate_status_transition(report,RegulatoryReport.Status.REVIEW_PENDING)
    report.report_status=RegulatoryReport.Status.REVIEW_PENDING; report.reviewed_by=user; report.save(); _audit(user,"REPORT_REVIEW_REQUEST",report); return report
@transaction.atomic
def approve_report(report,user,password=None,reason="",request=None):
    if user.role!="ADMIN": raise PermissionDenied("ADMIN만 승인할 수 있습니다.")
    _validate_status_transition(report,RegulatoryReport.Status.APPROVED)
    sign_approval(target=report,user=user,password=password,meaning="규제보고서 승인",reason=reason,request=request,forbid_self_user_id=report.created_by_id)
    before=snapshot(report); report.report_status=RegulatoryReport.Status.APPROVED; report.approved_by=user; report.save(); record_audit(user=user,action="REPORT_APPROVE",target=report,before=before,after=snapshot(report),reason=reason,request=request,require_reason=True); return report


@transaction.atomic
def update_report(report,user,**data):
    _allowed(user); before=snapshot(report)
    for key,value in data.items(): setattr(report,key,value)
    revoke_active_signatures(report,user=user,reason="승인된 중요 데이터 변경으로 재승인 필요")
    if report.report_status in {RegulatoryReport.Status.APPROVED,RegulatoryReport.Status.GENERATED}:
        report.report_status=RegulatoryReport.Status.DRAFT; report.approved_by=None
    report.save(); record_audit(user=user,action="REPORT_UPDATE",target=report,before=before,after=snapshot(report),reason="보고서 내용 수정"); return report
def _value(v): return str(v) if v not in (None,"") else "미입력"
def generate_docx_report(report,user,request=None):
    _allowed(user); validate_report_generation(report); _validate_status_transition(report,RegulatoryReport.Status.GENERATED); report.document_version+=1
    path=Path(settings.MEDIA_ROOT)/"reports"/f"{report.report_number}_v{report.document_version}.docx"; path.parent.mkdir(parents=True,exist_ok=True)
    doc=Document(); title=doc.add_heading(report.title,0); title.alignment=WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("본 문서는 포트폴리오 및 내부 업무 지원을 위한 예시 문서이며, 실제 규제기관 제출 전 담당자의 검토와 승인이 필요합니다.")
    sections=[("보고서 기본정보",[("보고서 번호",report.report_number),("규제기관",report.regulatory_authority),("문서 버전",report.document_version)]),("이상사례 기본정보",[("사건번호",report.adverse_event.event_number),("발생일",report.adverse_event.occurred_at.date()),("접수일",report.adverse_event.reported_at.date()),("심각도",report.adverse_event.severity),("보고 대상",report.adverse_event.reportability)]),("의료기기 및 추적성 정보",[("정보",report.device_information)]),("익명 환자 정보",[("정보",report.patient_information)]),("이상사례 상세 내용",[("요약",report.event_summary)]),("조사 내용",[("조사",report.investigation_summary),("근본 원인",report.root_cause_summary)]),("시정 및 예방조치",[("CAPA",report.capa_summary)]),("검토 및 승인 정보",[("검토자",report.reviewed_by),("승인자",report.approved_by)]),("결론",[("결론",report.conclusion)]),("작성 정보",[("작성자",report.created_by),("작성일",timezone.localdate())])]
    for heading,rows in sections:
        doc.add_heading(heading,level=1); table=doc.add_table(rows=0,cols=2)
        for k,v in rows: cells=table.add_row().cells; cells[0].text=k; cells[1].text=_value(v)
    doc.save(path); report.document_file=f"reports/{path.name}"; report.report_status=RegulatoryReport.Status.GENERATED; report.save(update_fields=["document_file","document_version","report_status","updated_at"]); _audit(user,"REPORT_DOCX",report,request=request); return path
def mark_report_submitted(report,user,request=None):
    _allowed(user)
    _validate_status_transition(report,RegulatoryReport.Status.SUBMITTED)
    report.report_status=RegulatoryReport.Status.SUBMITTED; report.submitted_at=timezone.now(); report.submitted_by=user; report.save(); _audit(user,"REPORT_SUBMIT",report,request=request); return report
def calculate_report_overdue_status(report): return report.is_overdue
